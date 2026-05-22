"""Generate DOCX export of a full conversation."""

import io
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def _add_label(doc: Document, text: str, color: RGBColor):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = color
    return p


def _add_body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Inches(0.2)
    return p


def build_docx(conversation: Dict[str, Any]) -> bytes:
    doc = Document()

    # Title
    title = conversation.get("title", "Conversation")
    _add_heading(doc, title, level=1)

    created_at = conversation.get("created_at", "")
    if created_at:
        p = doc.add_paragraph(f"Created: {created_at[:19].replace('T', ' ')}")
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    for msg in conversation.get("messages", []):
        if msg.get("role") == "user":
            _add_label(doc, "YOU", RGBColor(0x33, 0x33, 0x99))
            _add_body(doc, msg.get("content", ""))
            doc.add_paragraph()

        elif msg.get("role") == "assistant":
            _add_label(doc, "LLM COUNCIL", RGBColor(0x33, 0x66, 0x33))

            # Stage 1
            stage1: List[Dict] = msg.get("stage1") or []
            if stage1:
                _add_heading(doc, "Stage 1 — Individual Responses", level=2)
                for resp in stage1:
                    model = resp.get("model", "")
                    short = model.split("/")[1] if "/" in model else model
                    p = doc.add_paragraph()
                    run = p.add_run(short)
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                    p2 = doc.add_paragraph(resp.get("response", ""))
                    p2.paragraph_format.left_indent = Inches(0.2)
                    doc.add_paragraph()

            # Stage 2 — just aggregate rankings (skip raw eval wall of text)
            stage2: List[Dict] = msg.get("stage2") or []
            if stage2:
                _add_heading(doc, "Stage 2 — Peer Rankings", level=2)
                for rank in stage2:
                    model = rank.get("model", "")
                    short = model.split("/")[1] if "/" in model else model
                    parsed: List[str] = rank.get("parsed_ranking") or []
                    if parsed:
                        p = doc.add_paragraph()
                        p.add_run(f"{short} ranked: ").bold = True
                        p.add_run(", ".join(parsed))
                        p.paragraph_format.left_indent = Inches(0.2)
                doc.add_paragraph()

            # Stage 3
            stage3: Dict = msg.get("stage3") or {}
            if stage3:
                _add_heading(doc, "Stage 3 — Final Council Answer", level=2)
                chairman = stage3.get("model", "")
                short = chairman.split("/")[1] if "/" in chairman else chairman
                p = doc.add_paragraph()
                run = p.add_run(f"Chairman: {short}")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x2d, 0x8a, 0x2d)
                p2 = doc.add_paragraph(stage3.get("response", ""))
                p2.paragraph_format.left_indent = Inches(0.2)
                doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
